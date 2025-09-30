#!/usr/bin/env python3
"""
텍스트 추출 유틸리티 모듈
다양한 파일 형식에서 텍스트를 추출하고 청킹하는 기능 제공
"""

import io
import re
from typing import List, Union
from loguru import logger

try:
    import PyPDF2
    from docx import Document
    from bs4 import BeautifulSoup
except ImportError as e:
    logger.warning(f"선택적 의존성 누락: {e}")

def extract_from_txt(content: Union[bytes, str]) -> str:
    """TXT 파일에서 텍스트 추출"""
    try:
        if isinstance(content, bytes):
            text = content.decode('utf-8', errors='ignore')
        else:
            text = content
        
        # 기본 정리
        text = text.replace('\r\n', '\n').replace('\r', '\n')
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"TXT 텍스트 추출 오류: {e}")
        return ""

def extract_from_md(content: Union[bytes, str]) -> str:
    """Markdown 파일에서 텍스트 추출"""
    try:
        if isinstance(content, bytes):
            text = content.decode('utf-8', errors='ignore')
        else:
            text = content
        
        # Markdown 문법 간단히 정리 (완전한 파싱은 아님)
        # 헤더 마커 제거
        text = re.sub(r'^#{1,6}\s+', '', text, flags=re.MULTILINE)
        
        # 링크 텍스트만 추출: [텍스트](링크) -> 텍스트
        text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
        
        # 볼드, 이탤릭 마커 제거
        text = re.sub(r'\*{1,2}([^\*]+)\*{1,2}', r'\1', text)
        text = re.sub(r'_{1,2}([^_]+)_{1,2}', r'\1', text)
        
        # 코드 블록 마커 제거
        text = re.sub(r'```[^\n]*\n(.*?)\n```', r'\1', text, flags=re.DOTALL)
        text = re.sub(r'`([^`]+)`', r'\1', text)
        
        # 불필요한 공백 정리
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return text.strip()
        
    except Exception as e:
        logger.error(f"Markdown 텍스트 추출 오류: {e}")
        return ""

def extract_from_pdf(content: bytes) -> str:
    """PDF 파일에서 텍스트 추출"""
    try:
        pdf_file = io.BytesIO(content)
        pdf_reader = PyPDF2.PdfReader(pdf_file)
        
        text_parts = []
        
        for page_num, page in enumerate(pdf_reader.pages):
            try:
                page_text = page.extract_text()
                if page_text.strip():
                    text_parts.append(page_text)
                    
            except Exception as e:
                logger.warning(f"PDF 페이지 {page_num + 1} 추출 실패: {e}")
                continue
        
        full_text = '\n\n'.join(text_parts)
        
        # PDF 텍스트 정리
        # 불필요한 공백 제거
        full_text = re.sub(r'\s+', ' ', full_text)
        full_text = re.sub(r'\n\s*\n', '\n\n', full_text)
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"PDF 텍스트 추출 오류: {e}")
        return ""

def extract_from_docx(content: bytes) -> str:
    """DOCX 파일에서 텍스트 추출"""
    try:
        doc_file = io.BytesIO(content)
        doc = Document(doc_file)
        
        text_parts = []
        
        # 본문 텍스트 추출
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        
        # 표 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = []
                for cell in row.cells:
                    if cell.text.strip():
                        row_text.append(cell.text.strip())
                if row_text:
                    text_parts.append(' | '.join(row_text))
        
        full_text = '\n\n'.join(text_parts)
        
        # 불필요한 공백 정리
        full_text = re.sub(r'\s+', ' ', full_text)
        full_text = re.sub(r'\n\s*\n', '\n\n', full_text)
        
        return full_text.strip()
        
    except Exception as e:
        logger.error(f"DOCX 텍스트 추출 오류: {e}")
        return ""

def extract_from_html(content: Union[bytes, str]) -> str:
    """HTML에서 텍스트 추출"""
    try:
        if isinstance(content, bytes):
            html = content.decode('utf-8', errors='ignore')
        else:
            html = content
        
        soup = BeautifulSoup(html, 'html.parser')
        
        # 스크립트와 스타일 태그 제거
        for script in soup(["script", "style"]):
            script.decompose()
        
        # 텍스트 추출
        text = soup.get_text()
        
        # 불필요한 공백 정리
        lines = (line.strip() for line in text.splitlines())
        chunks = (phrase.strip() for line in lines for phrase in line.split("  "))
        text = ' '.join(chunk for chunk in chunks if chunk)
        
        return text
        
    except Exception as e:
        logger.error(f"HTML 텍스트 추출 오류: {e}")
        return ""

def chunk_text(
    text: str, 
    chunk_size: int = 1000, 
    overlap: int = 100,
    min_chunk_size: int = 50
) -> List[str]:
    """
    텍스트를 청크로 분할
    
    Args:
        text: 분할할 텍스트
        chunk_size: 청크 크기 (문자 수)
        overlap: 청크 간 겹치는 부분 크기
        min_chunk_size: 최소 청크 크기
    
    Returns:
        청크 리스트
    """
    if not text or len(text.strip()) < min_chunk_size:
        return []
    
    chunks = []
    start = 0
    text_length = len(text)
    
    while start < text_length:
        # 청크 끝 위치 계산
        end = start + chunk_size
        
        if end >= text_length:
            # 마지막 청크
            chunk = text[start:].strip()
            if len(chunk) >= min_chunk_size:
                chunks.append(chunk)
            break
        
        # 문장 경계에서 자르기 시도
        chunk_text = text[start:end]
        
        # 마지막 문장 끝 찾기
        sentence_endings = ['. ', '! ', '? ', '.\n', '!\n', '?\n']
        best_cut = -1
        
        for ending in sentence_endings:
            last_pos = chunk_text.rfind(ending)
            if last_pos > best_cut:
                best_cut = last_pos + len(ending)
        
        # 문장 끝을 찾지 못했으면 단어 경계에서 자르기
        if best_cut == -1:
            space_pos = chunk_text.rfind(' ')
            if space_pos > chunk_size * 0.7:  # 너무 짧지 않으면
                best_cut = space_pos + 1
            else:
                best_cut = end  # 강제로 자르기
        
        # 청크 추출
        chunk = text[start:start + best_cut].strip()
        
        if len(chunk) >= min_chunk_size:
            chunks.append(chunk)
        
        # 다음 시작 위치 (overlap 고려)
        start = start + best_cut - overlap
        
        # 무한 루프 방지
        if start <= 0:
            start = best_cut
    
    return chunks

def clean_text(text: str) -> str:
    """텍스트 기본 정리"""
    if not text:
        return ""
    
    # 연속된 공백을 하나로
    text = re.sub(r'\s+', ' ', text)
    
    # 연속된 줄바꿈을 최대 2개로
    text = re.sub(r'\n{3,}', '\n\n', text)
    
    # 앞뒤 공백 제거
    text = text.strip()
    
    return text

# 지원하는 파일 형식 매핑
EXTRACTORS = {
    'txt': extract_from_txt,
    'md': extract_from_md,
    'markdown': extract_from_md,
    'pdf': extract_from_pdf,
    'docx': extract_from_docx,
    'doc': extract_from_docx,
    'html': extract_from_html,
    'htm': extract_from_html
}

def extract_text_by_extension(content: Union[bytes, str], file_extension: str) -> str:
    """파일 확장자에 따라 적절한 추출기 선택"""
    file_extension = file_extension.lower().lstrip('.')
    
    if file_extension in EXTRACTORS:
        extractor = EXTRACTORS[file_extension]
        return extractor(content)
    else:
        logger.warning(f"지원하지 않는 파일 형식: {file_extension}")
        # 기본적으로 텍스트로 시도
        return extract_from_txt(content)

if __name__ == "__main__":
    # 테스트 코드
    test_text = """
    이것은 테스트 텍스트입니다. 여러 문장으로 구성되어 있습니다.
    
    두 번째 단락입니다. 청킹 테스트를 위해 충분히 긴 텍스트를 만들어보겠습니다.
    
    세 번째 단락도 추가해보겠습니다. 이렇게 하면 청킹이 잘 되는지 확인할 수 있을 것입니다.
    """
    
    chunks = chunk_text(test_text, chunk_size=100, overlap=20)
    
    print(f"원본 텍스트 길이: {len(test_text)}")
    print(f"청크 개수: {len(chunks)}")
    
    for i, chunk in enumerate(chunks):
        print(f"\n청크 {i+1} ({len(chunk)}자):")
        print(chunk)
